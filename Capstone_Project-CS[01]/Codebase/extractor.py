"""
Feature 1: Resume Data Extraction
Extracts structured information from unstructured resume documents (PDF/Word/TXT)
"""

import os
import json
import re
import pdfplumber
from docx import Document
from langchain_openai import ChatOpenAI
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import JsonOutputParser
from dotenv import load_dotenv

from exceptions import FileValidationError, DataValidationError, LLMResponseError
from validators import validate_existing_file, validate_non_empty_text
from logging_config import get_logger
from messages import print_error, error_file_operation

load_dotenv()

logger = get_logger(__name__)

EXTRACTOR_CHUNKING_THRESHOLD = int(os.getenv("EXTRACTOR_CHUNKING_THRESHOLD", "6000"))
EXTRACTOR_CHUNK_SIZE = int(os.getenv("EXTRACTOR_CHUNK_SIZE", "3500"))
EXTRACTOR_CHUNK_OVERLAP = int(os.getenv("EXTRACTOR_CHUNK_OVERLAP", "300"))
RESUME_SECTION_HEADERS = [
    "professional summary",
    "summary",
    "experience",
    "work experience",
    "skills",
    "education",
    "projects",
    "certifications",
    "achievements",
]


class ResumeExtractor:
    """Extract structured data from resume documents"""
    
    def __init__(self):
        """Initialize the extractor with LangChain LLM"""
        self.llm = ChatOpenAI(
            model=os.getenv("EXTRACTION_MODEL", "gpt-4o-mini"),
            temperature=float(os.getenv("EXTRACTION_TEMPERATURE", "0.3"))
        )
        
        # Define extraction prompt template
        self.extraction_prompt = ChatPromptTemplate.from_messages([
            ("system", """You are an expert resume parser. Extract structured information from the resume text.
            
Return a JSON object with the following structure:
{{
    "personal_details": {{
        "name": "Full Name",
        "email": "email@example.com",
        "phone": "phone number",
        "location": "city, state/country",
        "linkedin": "LinkedIn URL (if available)",
        "portfolio": "Portfolio/Website URL (if available)"
    }},
    "education": [
        {{
            "degree": "Degree name",
            "institution": "Institution name",
            "graduation_date": "Graduation date or expected date",
            "gpa": "GPA (if mentioned)",
            "honors": "Honors or distinctions (if any)"
        }}
    ],
    "work_experience": [
        {{
            "job_title": "Job title",
            "company": "Company name",
            "location": "Location",
            "start_date": "Start date",
            "end_date": "End date or 'Present'",
            "responsibilities": ["List of responsibilities and achievements"]
        }}
    ],
    "skills": {{
        "technical": ["List of technical skills"],
        "soft": ["List of soft skills"],
        "languages": ["Programming/spoken languages"],
        "tools": ["Tools and technologies"]
    }},
    "projects": [
        {{
            "name": "Project name",
            "description": "Brief description",
            "technologies": ["Technologies used"],
            "achievements": ["Key achievements or outcomes"]
        }}
    ],
    "certifications": [
        {{
            "name": "Certification name",
            "issuer": "Issuing organization",
            "date": "Date obtained",
            "credential_id": "ID (if available)"
        }}
    ],
    "achievements": [
        "List of notable achievements, awards, publications, or other accomplishments"
    ]
}}

If any field is not present in the resume, use null or an empty array as appropriate.
Be thorough and extract all relevant information."""),
            ("human", "{resume_text}")
        ])
        
        self.chain = self.extraction_prompt | self.llm
    
    def split_text_into_sections(self, text, section_headers=None):
        """Split resume text into semantic sections using common resume headings."""
        if not text or not isinstance(text, str):
            return []

        section_headers = section_headers or RESUME_SECTION_HEADERS
        normalized_headers = {header.lower() for header in section_headers}
        sections = []
        current_lines = []

        for line in text.splitlines():
            stripped = line.strip()
            normalized = stripped.rstrip(':').lower()

            if normalized in normalized_headers:
                if current_lines:
                    section_text = "\n".join(current_lines).strip()
                    if section_text:
                        sections.append(section_text)
                current_lines = [stripped]
                continue

            if stripped or current_lines:
                current_lines.append(line)

        if current_lines:
            section_text = "\n".join(current_lines).strip()
            if section_text:
                sections.append(section_text)

        return sections if len(sections) > 1 else []

    def chunk_text_by_size(self, text, max_chars=None, overlap_chars=None):
        """Split oversized text into bounded chunks with small overlap."""
        if not text or not isinstance(text, str):
            return []

        max_chars = max_chars or EXTRACTOR_CHUNK_SIZE
        overlap_chars = overlap_chars if overlap_chars is not None else EXTRACTOR_CHUNK_OVERLAP

        if len(text) <= max_chars:
            return [text.strip()]

        paragraphs = [paragraph.strip() for paragraph in re.split(r'\n\s*\n', text) if paragraph.strip()]
        if not paragraphs:
            paragraphs = [text.strip()]

        chunks = []
        current = ""

        for paragraph in paragraphs:
            candidate = f"{current}\n\n{paragraph}".strip() if current else paragraph
            if len(candidate) <= max_chars:
                current = candidate
                continue

            if current:
                chunks.append(current)
                overlap = current[-overlap_chars:].strip() if overlap_chars > 0 else ""
                current = f"{overlap}\n\n{paragraph}".strip() if overlap else paragraph
            else:
                start = 0
                while start < len(paragraph):
                    end = min(start + max_chars, len(paragraph))
                    chunk = paragraph[start:end].strip()
                    if chunk:
                        chunks.append(chunk)
                    if end >= len(paragraph):
                        current = ""
                        break
                    start = max(0, end - overlap_chars)

        if current:
            chunks.append(current)

        return chunks

    def maybe_chunk_document(self, text):
        """Chunk large resumes by section first, then by bounded size if needed."""
        if len(text) <= EXTRACTOR_CHUNKING_THRESHOLD:
            return [text]

        sections = self.split_text_into_sections(text)
        candidate_sections = sections or [text]
        chunks = []

        for section in candidate_sections:
            if len(section) <= EXTRACTOR_CHUNK_SIZE:
                chunks.append(section)
            else:
                chunks.extend(self.chunk_text_by_size(section))

        return chunks or [text]

    def _parse_llm_response(self, response):
        """Parse and validate a single LLM response payload."""
        content = response.content

        if not content:
            raise LLMResponseError(
                "❌ LLM returned empty response.",
                {"response": None}
            )

        if isinstance(content, str):
            if "```json" in content:
                content = content.split("```json")[1].split("```")[0].strip()
            elif "```" in content:
                content = content.split("```")[1].split("```")[0].strip()

            try:
                return json.loads(content)
            except json.JSONDecodeError as e:
                logger.error(f"JSON parse error: {str(e)}")
                logger.debug(f"Raw response (first 500 chars): {response.content[:500]}")
                raise LLMResponseError(
                    "❌ LLM returned invalid JSON. Please try again.",
                    {"json_error": str(e)}
                )

        return content

    def _merge_list_values(self, items):
        """Merge list values with order-preserving deduplication."""
        merged = []
        seen = set()

        for item in items or []:
            key = json.dumps(item, sort_keys=True, ensure_ascii=False) if isinstance(item, (dict, list)) else str(item)
            if key not in seen:
                seen.add(key)
                merged.append(item)

        return merged

    def merge_structured_data(self, chunk_results):
        """Merge chunked extraction results into a single structured profile."""
        merged = {
            "personal_details": {},
            "education": [],
            "work_experience": [],
            "skills": {
                "technical": [],
                "soft": [],
                "languages": [],
                "tools": [],
            },
            "projects": [],
            "certifications": [],
            "achievements": [],
        }

        for result in chunk_results:
            if not isinstance(result, dict):
                continue

            for key, value in (result.get("personal_details") or {}).items():
                if value and not merged["personal_details"].get(key):
                    merged["personal_details"][key] = value

            merged["education"] = self._merge_list_values(merged["education"] + (result.get("education") or []))
            merged["work_experience"] = self._merge_list_values(merged["work_experience"] + (result.get("work_experience") or []))
            merged["projects"] = self._merge_list_values(merged["projects"] + (result.get("projects") or []))
            merged["certifications"] = self._merge_list_values(merged["certifications"] + (result.get("certifications") or []))
            merged["achievements"] = self._merge_list_values(merged["achievements"] + (result.get("achievements") or []))

            result_skills = result.get("skills") or {}
            for skill_key in merged["skills"]:
                merged["skills"][skill_key] = self._merge_list_values(
                    merged["skills"][skill_key] + (result_skills.get(skill_key) or [])
                )

        return merged
    
    def extract_text_from_pdf(self, filepath):
        """Extract text from PDF file"""
        text = ""
        try:
            with pdfplumber.open(filepath) as pdf:
                if len(pdf.pages) == 0:
                    raise DataValidationError(
                        f"❌ PDF is empty: {filepath}",
                        {"path": filepath, "pages": 0}
                    )
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except pdfplumber.exceptions.PDFException as e:
            raise FileValidationError(
                f"❌ Invalid PDF file: {filepath}",
                {"path": filepath, "error": str(e)}
            )
        except Exception as e:
            raise FileValidationError(
                error_file_operation(filepath, "read", e),
                {"path": filepath, "error": str(e)}
            )
        return text
    
    def extract_text_from_docx(self, filepath):
        """Extract text from Word document"""
        text = ""
        try:
            doc = Document(filepath)
            if len(doc.paragraphs) == 0:
                logger.warning(f"DOCX document has no paragraphs: {filepath}")
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
        except FileNotFoundError:
            raise FileValidationError(
                f"❌ File not found: {filepath}",
                {"path": filepath}
            )
        except Exception as e:
            raise FileValidationError(
                error_file_operation(filepath, "read", e),
                {"path": filepath, "error": str(e)}
            )
        return text
    
    def extract_text_from_txt(self, filepath):
        """Extract text from TXT file with encoding fallback"""
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(filepath, 'r', encoding=encoding) as f:
                    text = f.read()
                logger.debug(f"Successfully read {filepath} with encoding {encoding}")
                return text
            except UnicodeDecodeError:
                logger.debug(f"Failed to read {filepath} with encoding {encoding}, trying next...")
                continue
            except IOError as e:
                raise FileValidationError(
                    f"❌ Cannot read file: {filepath}",
                    {"path": filepath, "error": str(e)}
                )
        
        # If all encodings failed
        raise FileValidationError(
            f"❌ File encoding error (expected UTF-8): {filepath}",
            {"path": filepath, "tried_encodings": encodings}
        )
    
    def convert_document_to_text(self, filepath):
        """
        Feature 1a: Document Conversion
        Convert PDF, Word, or TXT to plain text
        """
        try:
            validate_existing_file(filepath, "Resume file")
        except FileValidationError as e:
            logger.error(f"File validation failed: {e.message}")
            raise
        
        file_extension = os.path.splitext(filepath)[1].lower()
        
        print(f"📄 Converting document: {os.path.basename(filepath)}")
        
        if file_extension == '.pdf':
            return self.extract_text_from_pdf(filepath)
        elif file_extension in ['.docx', '.doc']:
            return self.extract_text_from_docx(filepath)
        elif file_extension == '.txt':
            return self.extract_text_from_txt(filepath)
        else:
            raise FileValidationError(
                f"❌ Unsupported file format: {file_extension}. Supported: PDF, DOCX, TXT",
                {"path": filepath, "extension": file_extension}
            )
    
    def extract_structured_data(self, resume_text):
        """
        Feature 1b: LLM-Based Extraction
        Use LangChain + GPT-4o-mini to extract structured JSON
        """
        try:
            validate_non_empty_text(resume_text, "Extracted resume text")
        except DataValidationError as e:
            logger.error(f"Resume text validation failed: {e.message}")
            raise
        
        print("🔍 Extracting structured data using LLM...")
        
        try:
            chunks = self.maybe_chunk_document(resume_text)
            if len(chunks) > 1:
                logger.info(f"Chunking large resume into {len(chunks)} chunks for extraction")
            
            chunk_results = []
            for chunk in chunks:
                response = self.chain.invoke({"resume_text": chunk})
                chunk_results.append(self._parse_llm_response(response))
            
            return chunk_results[0] if len(chunk_results) == 1 else self.merge_structured_data(chunk_results)
        except LLMResponseError:
            raise
        except Exception as e:
            logger.error(f"Extraction error: {str(e)}")
            raise LLMResponseError(
                f"❌ Error during extraction: {str(e)}",
                {"error": str(e)}
            )
    
    def save_to_json(self, data, output_path):
        """
        Feature 1c: JSON Output
        Save extracted data to JSON file
        """
        print(f"💾 Saving extracted data to: {output_path}")
        
        output_dir = os.path.dirname(output_path)
        if output_dir:
            os.makedirs(output_dir, exist_ok=True)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        
        print("✅ Data extraction complete!")
    
    def process_resume(self, input_filepath, output_filepath=None):
        """
        Complete Feature 1 workflow:
        1. Convert document to text
        2. Extract structured data using LLM
        3. Save to JSON
        """
        print("\n" + "="*60)
        print("FEATURE 1: RESUME DATA EXTRACTION")
        print("="*60 + "\n")
        
        try:
            # Step 1: Document conversion
            resume_text = self.convert_document_to_text(input_filepath)
            
            try:
                validate_non_empty_text(resume_text, "Extracted resume text")
            except DataValidationError as e:
                print(e.message)
                logger.error(f"Resume text validation failed: {e.message}")
                raise
            
            print(f"📝 Extracted {len(resume_text)} characters of text\n")
            
            # Step 2: LLM extraction
            structured_data = self.extract_structured_data(resume_text)
            
            # Step 3: Save to JSON
            if output_filepath is None:
                filename = os.path.splitext(os.path.basename(input_filepath))[0]
                output_filepath = os.path.join(
                    os.path.dirname(os.path.abspath(__file__)),
                    f"{filename}_extracted.json"
                )
            
            self.save_to_json(structured_data, output_filepath)
            
            return structured_data
        except (FileValidationError, DataValidationError, LLMResponseError) as e:
            print(e.message)
            logger.error(f"Resume processing failed: {e.message}", extra={"debug": e.debug_info})
            raise
        except Exception as e:
            error_msg = f"❌ Unexpected error during resume extraction: {str(e)}"
            print(error_msg)
            logger.error(error_msg, exc_info=True)
            raise


def main():
    """Test the extractor with a sample resume"""
    extractor = ResumeExtractor()
    
    # Example usage
    print("Resume Extractor - Feature 1")
    print("Supported formats: PDF, DOCX, TXT\n")
    
    resume_path = input("Enter the path to your resume file: ").strip()
    
    try:
        structured_data = extractor.process_resume(resume_path)
        print("\n✨ Extraction successful!")
        print(f"Name: {structured_data.get('personal_details', {}).get('name', 'N/A')}")
        print(f"Skills: {len(structured_data.get('skills', {}).get('technical', []))} technical skills found")
        print(f"Experience: {len(structured_data.get('work_experience', []))} positions found")
        
    except Exception as e:
        print(f"\n❌ Error: {str(e)}")


if __name__ == "__main__":
    main()
