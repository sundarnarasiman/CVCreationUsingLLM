"""
Output Formatters
Convert resume JSON to PDF and DOCX formats
"""

import os
import json
from datetime import datetime
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


class ResumePDFFormatter:
    """Generate PDF resume from JSON data"""
    
    def __init__(self):
        """Initialize PDF formatter"""
        self.pdf = None
    
    def create_pdf(self, resume_data, output_path):
        """Create a professional PDF resume"""
        try:
            self.pdf = FPDF('P', 'mm', 'A4')  # Explicitly set format
            self.pdf.set_left_margin(10)
            self.pdf.set_right_margin(10)
            self.pdf.add_page()
            self.pdf.set_auto_page_break(auto=True, margin=15)
            
            sections = resume_data.get('resume_sections', {})
            
            # Header
            try:
                self._add_header(sections.get('header', {}))
            except Exception as e:
                print(f"⚠️  Skipping header: {e}")
            
            # Professional Summary
            if 'professional_summary' in sections:
                try:
                    self._add_section_title("Professional Summary")
                    self._add_text(sections['professional_summary'], size=10)
                    self.pdf.ln(3)
                except Exception as e:
                    print(f"⚠️  Skipping professional summary: {e}")
                try:
                    self._add_skills_section(sections['skills'])
                except Exception as e:
                    print(f"⚠️  Skipping skills section: {e}")
            
            # Experience
            if 'experience' in sections:
                try:
                    self._add_experience_section(sections['experience'])
                except Exception as e:
                    print(f"⚠️  Skipping experience section: {e}")
            
            # Education
            if 'education' in sections:
                try:
                    self._add_education_section(sections['education'])
                except Exception as e:
                    print(f"⚠️  Skipping education section: {e}")
            
            # Projects
            if 'projects' in sections:
                try:
                    self._add_projects_section(sections['projects'])
                except Exception as e:
                    print(f"⚠️  Skipping projects section: {e}")
            
            # Certifications
            if 'certifications' in sections:
                try:
                    self._add_certifications_section(sections['certifications'])
                except Exception as e:
                    print(f"⚠️  Skipping certifications section: {e}")
            
            # Save PDF
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            self.pdf.output(output_path)
            print(f"📄 PDF resume saved to: {output_path}")
        except Exception as e:
            raise Exception(f"PDF creation failed: {str(e)}")
    
    def _add_header(self, header):
        """Add resume header with contact info"""
        self.pdf.set_font('Helvetica', 'B', 16)
        name = header.get('name', 'Your Name')
        self.pdf.cell(0, 10, name, ln=True, align='C')
        
        self.pdf.set_font('Helvetica', '', 9)
        contact_info = header.get('contact', [])
        contact_line = ' | '.join(contact_info[:4])
        self.pdf.cell(0, 5, contact_line, ln=True, align='C')
        self.pdf.ln(5)
    
    def _add_section_title(self, title):
        """Add section title"""
        self.pdf.set_font('Helvetica', 'B', 12)
        self.pdf.cell(0, 8, title.upper(), ln=True)
        self.pdf.ln(2)
    
    def _add_text(self, text, size=10, style=''):
        """Add regular text"""
        self.pdf.set_font('Helvetica', style, size)
        self.pdf.set_x(self.pdf.l_margin)
        self.pdf.multi_cell(0, 5, text)
    
    def _add_skills_section(self, skills):
        """Add skills section"""
        self._add_section_title("Skills")
        
        for skill_type, skill_list in skills.items():
            if skill_list:
                skill_type_formatted = skill_type.replace('_', ' ').title()
                skill_text = ', '.join(skill_list)
                full_text = f"{skill_type_formatted}: {skill_text}"
                
                self.pdf.set_font('Helvetica', '', 10)
                self.pdf.set_x(self.pdf.l_margin)
                self.pdf.multi_cell(0, 5, full_text)
        
        self.pdf.ln(2)
        
        self.pdf.ln(2)
    
    def _add_experience_section(self, experiences):
        """Add experience section"""
        self._add_section_title("Professional Experience")
        
        for exp in experiences:
            # Job title and company
            self.pdf.set_font('Helvetica', 'B', 11)
            title_company = f"{exp.get('title', 'Position')} | {exp.get('company', 'Company')}"
            self.pdf.cell(0, 6, title_company, ln=True)
            
            # Location and dates
            self.pdf.set_font('Helvetica', 'I', 9)
            location_dates = f"{exp.get('location', '')} | {exp.get('dates', '')}"
            self.pdf.cell(0, 5, location_dates, ln=True)
            
            # Bullets
            self.pdf.set_font('Helvetica', '', 10)
            bullets = exp.get('bullets', [])
            for bullet in bullets:
                # Add bullet with indentation
                indent_text = "    - " + bullet  # 4 spaces + dash for indentation
                self.pdf.set_x(self.pdf.l_margin)
                self.pdf.multi_cell(0, 5, indent_text)
            
            self.pdf.ln(2)
    
    def _add_education_section(self, education):
        """Add education section"""
        self._add_section_title("Education")
        
        for edu in education:
            self.pdf.set_font('Helvetica', 'B', 10)
            degree_institution = f"{edu.get('degree', 'Degree')} | {edu.get('institution', 'Institution')}"
            self.pdf.cell(0, 6, degree_institution, ln=True)
            
            self.pdf.set_font('Helvetica', 'I', 9)
            grad_date = edu.get('graduation_date', '')
            honors = edu.get('honors', '')
            if honors:
                self.pdf.cell(0, 5, f"{grad_date} | {honors}", ln=True)
            else:
                self.pdf.cell(0, 5, grad_date, ln=True)
            
            self.pdf.ln(1)
    
    def _add_projects_section(self, projects):
        """Add projects section"""
        self._add_section_title("Projects")
        
        for proj in projects:
            self.pdf.set_font('Helvetica', 'B', 10)
            self.pdf.cell(0, 6, proj.get('name', 'Project'), ln=True)
            
            self.pdf.set_font('Helvetica', '', 10)
            description = proj.get('description', '')
            self.pdf.set_x(self.pdf.l_margin)
            self.pdf.multi_cell(0, 5, description)
            
            technologies = proj.get('technologies', [])
            if technologies:
                self.pdf.set_font('Helvetica', 'I', 9)
                tech_text = f"Technologies: {', '.join(technologies)}"
                self.pdf.set_x(self.pdf.l_margin)
                self.pdf.multi_cell(0, 5, tech_text)
            
            self.pdf.ln(1)
    
    def _add_certifications_section(self, certifications):
        """Add certifications section"""
        self._add_section_title("Certifications")
        
        for cert in certifications:
            self.pdf.set_font('Helvetica', '', 10)
            cert_text = f"- {cert.get('name', 'Certification')} - {cert.get('issuer', '')} ({cert.get('date', '')})"
            self.pdf.set_x(self.pdf.l_margin)
            self.pdf.multi_cell(0, 5, cert_text)


class ResumeDOCXFormatter:
    """Generate DOCX resume from JSON data"""
    
    def __init__(self):
        """Initialize DOCX formatter"""
        self.doc = None
    
    def create_docx(self, resume_data, output_path):
        """Create a professional DOCX resume"""
        self.doc = Document()
        
        # Set document margins
        sections = resume_data.get('resume_sections', {})
        
        # Header
        self._add_header(sections.get('header', {}))
        
        # Professional Summary
        if 'professional_summary' in sections:
            self._add_section_title("PROFESSIONAL SUMMARY")
            self._add_paragraph(sections['professional_summary'])
        
        # Skills
        if 'skills' in sections:
            self._add_skills_section(sections['skills'])
        
        # Experience
        if 'experience' in sections:
            self._add_experience_section(sections['experience'])
        
        # Education
        if 'education' in sections:
            self._add_education_section(sections['education'])
        
        # Projects
        if 'projects' in sections:
            self._add_projects_section(sections['projects'])
        
        # Certifications
        if 'certifications' in sections:
            self._add_certifications_section(sections['certifications'])
        
        # Save DOCX
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        self.doc.save(output_path)
        print(f"📄 DOCX resume saved to: {output_path}")
    
    def _add_header(self, header):
        """Add resume header with contact info"""
        # Name
        name_para = self.doc.add_paragraph()
        name_run = name_para.add_run(header.get('name', 'Your Name'))
        name_run.font.size = Pt(16)
        name_run.font.bold = True
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contact info
        contact_info = header.get('contact', [])
        contact_para = self.doc.add_paragraph(' | '.join(contact_info[:4]))
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_para.runs[0].font.size = Pt(10)
        
        self.doc.add_paragraph()  # Spacing
    
    def _add_section_title(self, title):
        """Add section title"""
        para = self.doc.add_paragraph()
        run = para.add_run(title)
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Add bottom border
        para.paragraph_format.border_bottom.width = Pt(1)
    
    def _add_paragraph(self, text, size=10, bold=False, italic=False):
        """Add a paragraph"""
        para = self.doc.add_paragraph(text)
        para.runs[0].font.size = Pt(size)
        para.runs[0].font.bold = bold
        para.runs[0].font.italic = italic
        return para
    
    def _add_skills_section(self, skills):
        """Add skills section"""
        self._add_section_title("SKILLS")
        
        for skill_type, skill_list in skills.items():
            if skill_list:
                skill_type_formatted = skill_type.replace('_', ' ').title()
                para = self.doc.add_paragraph()
                
                type_run = para.add_run(f"{skill_type_formatted}: ")
                type_run.font.bold = True
                type_run.font.size = Pt(10)
                
                skills_run = para.add_run(', '.join(skill_list))
                skills_run.font.size = Pt(10)
        
        self.doc.add_paragraph()  # Spacing
    
    def _add_experience_section(self, experiences):
        """Add experience section"""
        self._add_section_title("PROFESSIONAL EXPERIENCE")
        
        for exp in experiences:
            # Title and company
            title_para = self.doc.add_paragraph()
            title_run = title_para.add_run(f"{exp.get('title', 'Position')} | {exp.get('company', 'Company')}")
            title_run.font.size = Pt(11)
            title_run.font.bold = True
            
            # Location and dates
            loc_para = self.doc.add_paragraph(f"{exp.get('location', '')} | {exp.get('dates', '')}")
            loc_para.runs[0].font.size = Pt(9)
            loc_para.runs[0].font.italic = True
            
            # Bullets
            bullets = exp.get('bullets', [])
            for bullet in bullets:
                bullet_para = self.doc.add_paragraph(bullet, style='List Bullet')
                bullet_para.runs[0].font.size = Pt(10)
            
            self.doc.add_paragraph()  # Spacing
    
    def _add_education_section(self, education):
        """Add education section"""
        self._add_section_title("EDUCATION")
        
        for edu in education:
            # Degree and institution
            deg_para = self.doc.add_paragraph()
            deg_run = deg_para.add_run(f"{edu.get('degree', 'Degree')} | {edu.get('institution', 'Institution')}")
            deg_run.font.size = Pt(10)
            deg_run.font.bold = True
            
            # Date and honors
            grad_date = edu.get('graduation_date', '')
            honors = edu.get('honors', '')
            info_text = f"{grad_date} | {honors}" if honors else grad_date
            
            info_para = self.doc.add_paragraph(info_text)
            info_para.runs[0].font.size = Pt(9)
            info_para.runs[0].font.italic = True
        
        self.doc.add_paragraph()  # Spacing
    
    def _add_projects_section(self, projects):
        """Add projects section"""
        self._add_section_title("PROJECTS")
        
        for proj in projects:
            # Project name
            name_para = self.doc.add_paragraph()
            name_run = name_para.add_run(proj.get('name', 'Project'))
            name_run.font.size = Pt(10)
            name_run.font.bold = True
            
            # Description
            desc_para = self.doc.add_paragraph(proj.get('description', ''))
            desc_para.runs[0].font.size = Pt(10)
            
            # Technologies
            technologies = proj.get('technologies', [])
            if technologies:
                tech_para = self.doc.add_paragraph(f"Technologies: {', '.join(technologies)}")
                tech_para.runs[0].font.size = Pt(9)
                tech_para.runs[0].font.italic = True
        
        self.doc.add_paragraph()  # Spacing
    
    def _add_certifications_section(self, certifications):
        """Add certifications section"""
        self._add_section_title("CERTIFICATIONS")
        
        for cert in certifications:
            cert_text = f"{cert.get('name', 'Certification')} - {cert.get('issuer', '')} ({cert.get('date', '')})"
            cert_para = self.doc.add_paragraph(cert_text, style='List Bullet')
            cert_para.runs[0].font.size = Pt(10)


def format_resume(resume_json_path, output_format='both'):
    """
    Format resume from JSON to PDF and/or DOCX
    
    Args:
        resume_json_path: Path to resume JSON file
        output_format: 'pdf', 'docx', or 'both'
    """
    # Load resume data
    with open(resume_json_path, 'r', encoding='utf-8') as f:
        resume_data = json.load(f)
    
    # Generate output paths
    base_path = os.path.splitext(resume_json_path)[0]
    pdf_path = f"{base_path}.pdf"
    docx_path = f"{base_path}.docx"
    
    print("\n" + "="*60)
    print("RESUME FORMATTING")
    print("="*60 + "\n")
    
    # Generate PDF
    if output_format in ['pdf', 'both']:
        try:
            pdf_formatter = ResumePDFFormatter()
            pdf_formatter.create_pdf(resume_data, pdf_path)
        except Exception as e:
            print(f"❌ Error generating PDF: {str(e)}")
    
    # Generate DOCX
    if output_format in ['docx', 'both']:
        try:
            docx_formatter = ResumeDOCXFormatter()
            docx_formatter.create_docx(resume_data, docx_path)
        except Exception as e:
            print(f"❌ Error generating DOCX: {str(e)}")
    
    print("\n✅ Resume formatting complete!")
    
    return pdf_path, docx_path


def main():
    """Test the formatters"""
    print("Resume Formatters")
    print("="*60)
    
    resume_path = input("\nEnter path to resume JSON: ").strip()
    
    print("\nOutput format:")
    print("  1. PDF only")
    print("  2. DOCX only")
    print("  3. Both PDF and DOCX")
    
    choice = input("\nSelect (1/2/3): ").strip()
    
    format_map = {'1': 'pdf', '2': 'docx', '3': 'both'}
    output_format = format_map.get(choice, 'both')
    
    try:
        pdf_path, docx_path = format_resume(resume_path, output_format)
        print("\n✨ Formatting successful!")
        
    except Exception as e:
        print(f"\n❌ Error: {str(e)}")


if __name__ == "__main__":
    main()
