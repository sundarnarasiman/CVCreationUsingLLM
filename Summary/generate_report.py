from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Preformatted
from reportlab.lib import colors

# Cross-platform font configuration
DOCX_FONT = "Calibri"  # Available on all platforms via Office/LibreOffice
PDF_FONT = "Helvetica"  # Universal font available on all platforms

# Full report content
report_content = {
    "title": "CV CREATION USING DUAL-LLM SYSTEM\nProject Report",
    "abstract": "This capstone project presents a dual-LLM system for generating tailored, ATS-optimized CVs using OpenAI's GPT-4o-mini and GPT-4o models with LangChain orchestration. The system automates resume extraction from unstructured documents, parses job requirements, strategically aligns candidate profiles, and generates optimized resumes with real-time ATS feedback. The implementation includes semantic similarity scoring using OpenAI embeddings (text-embedding-3-small), frequency-weighted (TF) keyword selection via Python Counter, curated n-gram phrase detection for 17 multi-word technical terms, and threshold-based section-aware document chunking with structured merge logic — ensuring robustness at scale. This report captures architecture decisions and engineering trade-offs across model selection, chunking strategy, embedding choice, and regex-vs-semantic matching design.",
    
    "section1": {
        "title": "1. Introduction",
        "content": "Context: Most candidates fail to effectively tailor resumes for specific job postings, resulting in ATS rejection. Traditional resume optimization is manual, expensive, and time-consuming.\n\nMotivation: This project leverages dual-LLM architecture to optimize cost-performance trade-offs while delivering professional-grade resume optimization accessible to all job seekers. The combination of extraction complexity, NLP integration, and practical impact makes this an ideal capstone project demonstrating advanced AI and software engineering skills."
    },
    
    "section2": {
        "title": "2. Problem Statement",
        "content": "Core Issues:\n• Resumes fail ATS screening due to missing keywords and poor formatting\n• Candidates struggle to identify transferable skills relevant to target roles\n• Manual resume tailoring for multiple applications is time-consuming\n• Resume quality varies widely based on individual expertise\n\nMarket Gap: Existing tools are limited, expensive, inflexible, or low-quality. This project provides an intelligent, accessible, affordable solution."
    },
    
    "section3": {
        "title": "3. Objectives",
        "objectives": [
            "Automated resume extraction from PDF/DOCX/TXT",
            "Intelligent job description parsing",
            "Strategic resume generation (2-step process)",
            "Iterative revision with ATS feedback",
            "ATS scoring system (0-100 scale)",
            "Semantic keyword matching using embeddings",
            "Multi-format export (PDF/DOCX)",
            "Modular, scalable architecture",
            "Profile-job matching assessment",
            "Intuitive menu-driven interface",
            "Cost-optimized LLM orchestration"
        ]
    },
    
    "section4": {
        "title": "4. Methodology",
        "tech_stack": [
            "LangChain >=1.2.13 (with langchain-community and langchain-openai)",
            "GPT-4o-mini (extraction), GPT-4o (generation)",
            "OpenAI text-embedding-3-small (semantic similarity)",
            "pdfplumber, python-docx, fpdf2, reportlab",
            "Python 3.8+"
        ],
        "modules": [
            ("extractor.py", "Resume extraction", "Multi-format support"),
            ("parser.py", "Job parsing", "Keyword identification"),
            ("generator.py", "Resume generation", "Two-step process"),
            ("reviser.py", "Iterative refinement", "User edit integration"),
            ("ats_checker.py", "ATS scoring", "Semantic keyword evaluation (0-100 scale)"),
            ("matcher.py", "Profile-job matching", "Embedding-based gap analysis"),
            ("formatters.py", "PDF/DOCX export", "Professional formatting"),
            ("main.py", "Orchestration", "Complete workflow")
        ]
    },

    "section5": {
        "title": "5. System Architecture",
        "overview": "The system follows a pipeline-plus-feedback architecture. Inputs are parsed into structured JSON, scored for fit and ATS compatibility, then iteratively refined before final export.",
        "diagram": [
            "+----------------------+      +----------------------+",
            "| Resume (PDF/DOCX/TXT)|      | Job Description (TXT)|",
            "+----------+-----------+      +----------+-----------+",
            "           |                             |",
            "     extractor.py                   parser.py",
            "           +-------------+  +------------+",
            "                         v  v",
            "                    main.py (workflow)",
            "                         |",
            "                     matcher.py",
            "                         |",
            "                    generator.py",
            "                         |",
            "           +-------------+-------------+",
            "           |                           |",
            "     ats_checker.py <---- reviser.py (loop)",
            "           |",
            "      formatters.py",
            "           |",
            "+----------v---------------------------+",
            "| Tailored Resume Output (PDF/DOCX/JSON)|",
            "+---------------------------------------+"
        ],
        "flow": [
            "extractor.py converts unstructured resume content into normalized profile JSON.",
            "parser.py structures the target job description into requirements and keywords.",
            "main.py orchestrates feature execution and handles the interactive menu flow.",
            "matcher.py computes profile-to-job alignment before generation.",
            "generator.py runs strategic analysis then tailored resume synthesis.",
            "ats_checker.py computes ATS score and semantic keyword coverage.",
            "reviser.py applies user edits and re-runs ATS checks in a feedback loop.",
            "formatters.py exports final resumes to human-readable PDF and DOCX." 
        ]
    },

    "section6": {
        "title": "6. LLM Model Selection Trade-offs (Pros & Cons)",
        "rows": [
            {
                "option": "GPT-4o-mini",
                "pros": "Lower cost; lower latency; good for extraction and parsing tasks.",
                "cons": "Lower generation depth; weaker on nuanced phrasing for final resumes.",
                "best_for": "High-volume structured preprocessing"
            },
            {
                "option": "GPT-4o",
                "pros": "Higher quality synthesis; stronger reasoning for strategic tailoring.",
                "cons": "Higher cost; slower response time compared to mini model.",
                "best_for": "Final resume generation and refinement"
            },
            {
                "option": "Dual-Model (Current)",
                "pros": "Balanced cost/performance; quality where needed, efficiency elsewhere.",
                "cons": "More orchestration complexity and model-boundary design overhead.",
                "best_for": "Production pipeline with quality and budget constraints"
            }
        ]
    },

    "section7": {
        "title": "7. Chunking Strategy Trade-offs & Implementation",
        "rows": [
            {
                "option": "No Chunking (Previous)",
                "pros": "Preserves full context in one pass for short documents.",
                "cons": "Fails on long inputs due to token limits; no merge logic needed but no protection either.",
                "best_for": "Short resumes or compact job descriptions only"
            },
            {
                "option": "Fixed-Size Chunking",
                "pros": "Simple implementation; predictable token usage and latency.",
                "cons": "Can split related context across chunk boundaries, reducing extraction quality.",
                "best_for": "Baseline scalable processing where section fidelity is low priority"
            },
            {
                "option": "Section-Aware + Bounded (Current)",
                "pros": "Preserves section context; falls back to size chunking; char-threshold is env-tunable; per-chunk LLM calls with structured merge.",
                "cons": "More implementation complexity; merge logic must handle deduplication consistently.",
                "best_for": "Production extraction pipelines — implemented in extractor.py and parser.py"
            }
        ]
    },
    "section7b": {
        "title": "7b. Chunking Implementation Details",
        "thresholds": [
            ("extractor.py", "6,000 chars", "3,500 chars", "300 chars"),
            ("parser.py",   "5,000 chars", "3,000 chars", "250 chars")
        ],
        "stages": [
            "Stage 0: If len(text) <= threshold, send whole document — no overhead.",
            "Stage 1: Split by known section headers (e.g. EXPERIENCE, SKILLS, REQUIREMENTS).",
            "Stage 2: If any section exceeds chunk_size, apply paragraph-aware size splitting with overlap.",
            "Merge: Scalar fields keep first non-null value; list fields use order-preserving JSON-key deduplication; nested skill/keyword dicts merged per sub-category."
        ]
    },
    "section7c": {
        "title": "7c. Keyword Optimisation — Priorities 1 & 2",
        "rows": [
            {
                "option": "Priority 1: TF Weighting",
                "pros": "Most frequent / important keywords surface first; skills and technologies double-weighted.",
                "cons": "Frequency alone does not capture domain rarity (TF-IDF not yet applied).",
                "best_for": "Top-30 keyword selection in matcher.py and ats_checker.py"
            },
            {
                "option": "Priority 2: N-Gram Phrases",
                "pros": "17 curated technical phrases (e.g. machine learning, ci/cd); frequency-preserving via repeated list entries; no arbitrary bigram noise.",
                "cons": "Phrase list is manually maintained; new terminology requires explicit addition.",
                "best_for": "High-precision multi-word skill detection in both modules"
            }
        ]
    },

    "section8": {
        "title": "8. Embedding Strategy Trade-offs (Pros & Cons)",
        "rows": [
            {
                "option": "text-embedding-3-small (Current)",
                "pros": "Good semantic quality at low cost; fast enough for interactive use.",
                "cons": "May miss subtle domain-specific nuance vs larger models.",
                "best_for": "General semantic keyword matching at scale"
            },
            {
                "option": "Larger Embedding Models",
                "pros": "Potentially better semantic precision for complex terminology.",
                "cons": "Higher cost and latency; larger vector/storage footprint.",
                "best_for": "High-stakes matching where precision is critical"
            },
            {
                "option": "No Embeddings (Exact Overlap)",
                "pros": "Fully explainable logic; no embedding API dependency.",
                "cons": "Poor recall for related terms and synonyms.",
                "best_for": "Strict deterministic keyword filtering"
            }
        ]
    },

    "section9": {
        "title": "9. Regex vs Semantic Comparison Trade-offs (Pros & Cons)",
        "rows": [
            {
                "option": "Regex/Pattern Matching",
                "pros": "High interpretability; deterministic outputs; easy debugging.",
                "cons": "Lower recall for paraphrases and semantic variants.",
                "best_for": "Rule-driven compliance and strict keyword checks"
            },
            {
                "option": "Semantic Embedding Matching",
                "pros": "Higher recall across synonyms and related skills.",
                "cons": "Harder to explain to end users; threshold tuning required.",
                "best_for": "Robust skill similarity and ATS relevance scoring"
            },
            {
                "option": "Hybrid (Regex + Semantic)",
                "pros": "Balances precision and recall with layered safeguards.",
                "cons": "Increased maintenance complexity and orchestration effort.",
                "best_for": "Production systems requiring trust + flexibility"
            }
        ]
    },

    "section10": {
        "title": "10. Testing and Validation Summary",
        "validation_points": [
            "53/53 automated tests passing across all modules (unittest framework).",
            "Chunking helpers tested: section splitting, size-bounded chunking, small-document bypass.",
            "Merge logic tested: scalar first-wins, list deduplication, nested skill/keyword dicts.",
            "Multi-chunk LLM flows tested with content-aware mocks in extractor and parser.",
            "N-gram phrase detection verified: curated phrases ranked above single tokens.",
            "Frequency-weighted (TF) keyword selection verified: most_common(30) returns highest-frequency terms.",
            "End-to-end workflow verified from resume input to PDF/DOCX output.",
            "Semantic matching behavior validated with related-skill scenarios.",
            "ATS scoring feedback loop tested during iterative resume revision.",
            "Error handling pathways tested for missing inputs, empty files, and invalid JSON responses."
        ],
        "quality_metrics": [
            "Test suite: 53 tests, 0 failures, 0 errors, 0 skipped.",
            "Resume extraction: robust on structured and semi-structured resumes; auto-chunks large inputs.",
            "Keyword ranking: frequency-weighted top-30 selection improves ATS precision over position-based selection.",
            "Multi-word phrase matching: 17 technical phrases detected and frequency-weighted correctly.",
            "Semantic relevance: improved matching quality over exact keyword overlap baselines.",
            "Maintainability: modular design supports isolated updates by component."
        ]
    },
    
    "section11": {
        "title": "11. Results and Analysis",
        "status": "FULLY COMPLETE AND OPERATIONAL — ALL THREE CHUNKING PRIORITIES IMPLEMENTED",
        "results": [
            "8 fully functional core modules with modular execution",
            "4 core features with supporting ATS and matching utilities",
            "Production-ready deployment with comprehensive error handling",
            "Priority 1: Frequency-weighted (TF) keyword selection — Counter.most_common(30) in matcher.py and ats_checker.py",
            "Priority 2: Curated n-gram phrase detection — 17 technical phrases with frequency preservation",
            "Priority 3: Threshold-based section-aware document chunking with structured merge in extractor.py and parser.py",
            "Semantic similarity matching integrated for profile-job matching and ATS keyword analysis",
            "53/53 automated tests passing including dedicated chunking and merge coverage",
            "Cross-platform report and resume export support (PDF/DOCX)"
        ]
    },
    
    "section12": {
        "title": "12. Conclusion",
        "contributions": [
            "Fully automated resume optimization pipeline",
            "Intelligent job-profile alignment with semantic matching",
            "ATS-optimized resume generation (80%+ compatibility)",
            "Real-time feedback and revision capability",
            "Cost-efficient dual-LLM orchestration",
            "Frequency-weighted keyword selection and curated n-gram phrase detection",
            "Threshold-based document chunking with section-aware splitting and structured merge",
            "Comprehensive 53-test suite covering edge cases, error handling, and chunking logic",
            "Modular, maintainable architecture with environment-variable tunable chunking thresholds"
        ],
        "future": [
            "Multi-language support (6 languages)",
            "Cloud deployment (REST API, web app)",
            "Advanced analytics dashboard",
            "Industry-specific templates",
            "Integration with LinkedIn, job boards, ATS platforms"
        ]
    }
}


def set_cell_font(cell, font_name=DOCX_FONT, font_size=10, bold=False):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.font.bold = bold


def add_docx_bullets(doc, items, style='List Bullet'):
    for item in items:
        para = doc.add_paragraph(item, style=style)
        for run in para.runs:
            run.font.name = DOCX_FONT
            run.font.size = Pt(10)


def add_docx_tradeoff_table(doc, section):
    heading = doc.add_heading(section["title"], level=2)
    heading.runs[0].font.name = DOCX_FONT

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    headers = ["Option", "Pros", "Cons", "Best Fit"]
    for idx, text in enumerate(headers):
        header_cells[idx].text = text
        set_cell_font(header_cells[idx], font_name=DOCX_FONT, font_size=10, bold=True)

    for row in section["rows"]:
        cells = table.add_row().cells
        cells[0].text = row["option"]
        cells[1].text = row["pros"]
        cells[2].text = row["cons"]
        cells[3].text = row["best_for"]
        for cell in cells:
            set_cell_font(cell, font_name=DOCX_FONT, font_size=10)

    doc.add_paragraph()


def add_pdf_bullets(story, items, body_style):
    for item in items:
        story.append(Paragraph(f"<bullet>•</bullet> {item}", body_style))


def add_pdf_tradeoff_table(story, section, body_style):
    story.append(Paragraph(section["title"], body_style))

    table_data = [["Option", "Pros", "Cons", "Best Fit"]]
    for row in section["rows"]:
        table_data.append([
            Paragraph(row["option"], body_style),
            Paragraph(row["pros"], body_style),
            Paragraph(row["cons"], body_style),
            Paragraph(row["best_for"], body_style)
        ])

    table = Table(table_data, colWidths=[1.45 * inch, 2.0 * inch, 2.0 * inch, 1.35 * inch])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#D9E2F3')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('FONTNAME', (0, 0), (-1, 0), PDF_FONT),
        ('FONTSIZE', (0, 0), (-1, -1), 8.5),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
    ]))
    story.append(table)
    story.append(Spacer(1, 0.1 * inch))

# ============ DOCX Generation ============
def generate_docx():
    doc = Document()
    
    # Set default font for document
    style = doc.styles['Normal']
    style.font.name = DOCX_FONT
    style.font.size = Pt(11)
    
    # Title
    title = doc.add_paragraph(report_content["title"])
    title_format = title.paragraph_format
    title_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.name = DOCX_FONT
    title_run.font.size = Pt(14)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    
    doc.add_paragraph()  # Spacing
    
    # Abstract
    abstract_heading = doc.add_heading("Abstract", level=2)
    abstract_heading_run = abstract_heading.runs[0]
    abstract_heading_run.font.name = DOCX_FONT
    
    abstract_para = doc.add_paragraph(report_content["abstract"])
    for run in abstract_para.runs:
        run.font.name = DOCX_FONT
        run.font.size = Pt(10)
    
    doc.add_page_break()
    
    # Sections 1-2
    for i in [1, 2]:
        section_key = f"section{i}"
        section_heading = doc.add_heading(report_content[section_key]["title"], level=2)
        section_heading_run = section_heading.runs[0]
        section_heading_run.font.name = DOCX_FONT
        
        section_para = doc.add_paragraph(report_content[section_key]["content"])
        for run in section_para.runs:
            run.font.name = DOCX_FONT
            run.font.size = Pt(11)
    
    doc.add_page_break()
    
    # Section 3
    sec3_heading = doc.add_heading(report_content["section3"]["title"], level=2)
    sec3_heading.runs[0].font.name = DOCX_FONT
    add_docx_bullets(doc, report_content["section3"]["objectives"])
    
    # Section 4
    sec4_heading = doc.add_heading(report_content["section4"]["title"], level=2)
    sec4_heading.runs[0].font.name = DOCX_FONT
    
    tech_label = doc.add_paragraph("Technology Stack:", style='List Bullet')
    tech_label.runs[0].font.name = DOCX_FONT
    tech_label.runs[0].font.bold = True
    
    for tech in report_content["section4"]["tech_stack"]:
        tech_para = doc.add_paragraph(tech, style='List Bullet 2')
        for run in tech_para.runs:
            run.font.name = DOCX_FONT
    
    modules_label = doc.add_paragraph("\nCore Modules:")
    modules_label.runs[0].font.name = DOCX_FONT
    modules_label.runs[0].font.bold = True
    
    for module, responsibility, feature in report_content["section4"]["modules"]:
        mod_para = doc.add_paragraph(f"{module}: {responsibility} ({feature})", style='List Bullet')
        for run in mod_para.runs:
            run.font.name = DOCX_FONT
            run.font.size = Pt(10)

    doc.add_page_break()

    # Section 5 - Architecture
    sec5_heading = doc.add_heading(report_content["section5"]["title"], level=2)
    sec5_heading.runs[0].font.name = DOCX_FONT

    overview_para = doc.add_paragraph(report_content["section5"]["overview"])
    for run in overview_para.runs:
        run.font.name = DOCX_FONT
        run.font.size = Pt(10)

    diag_label = doc.add_paragraph("Architecture Diagram:")
    diag_label.runs[0].font.name = DOCX_FONT
    diag_label.runs[0].font.bold = True

    for line in report_content["section5"]["diagram"]:
        line_para = doc.add_paragraph(line)
        for run in line_para.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(8)

    flow_label = doc.add_paragraph("Component Interaction Notes:")
    flow_label.runs[0].font.name = DOCX_FONT
    flow_label.runs[0].font.bold = True
    add_docx_bullets(doc, report_content["section5"]["flow"])

    doc.add_page_break()

    # Sections 6-9 Trade-offs
    for section_key in ["section6", "section7", "section8", "section9"]:
        add_docx_tradeoff_table(doc, report_content[section_key])

    # Section 7b – Chunking Implementation Details
    sec7b = report_content["section7b"]
    sec7b_heading = doc.add_heading(sec7b["title"], level=2)
    sec7b_heading.runs[0].font.name = DOCX_FONT

    thresh_label = doc.add_paragraph("Threshold Configuration:")
    thresh_label.runs[0].font.name = DOCX_FONT
    thresh_label.runs[0].font.bold = True

    thresh_table = doc.add_table(rows=1, cols=4)
    thresh_table.style = 'Table Grid'
    hdr = thresh_table.rows[0].cells
    for idx, text in enumerate(["Module", "Doc Threshold", "Chunk Size", "Overlap"]):
        hdr[idx].text = text
        set_cell_font(hdr[idx], font_name=DOCX_FONT, font_size=10, bold=True)
    for row_data in sec7b["thresholds"]:
        cells = thresh_table.add_row().cells
        for idx, val in enumerate(row_data):
            cells[idx].text = val
            set_cell_font(cells[idx], font_name=DOCX_FONT, font_size=10)
    doc.add_paragraph()

    stages_label = doc.add_paragraph("Chunking Stages:")
    stages_label.runs[0].font.name = DOCX_FONT
    stages_label.runs[0].font.bold = True
    add_docx_bullets(doc, sec7b["stages"])

    # Section 7c – Keyword Optimisation Trade-offs
    add_docx_tradeoff_table(doc, report_content["section7c"])

    # Section 10
    sec10_heading = doc.add_heading(report_content["section10"]["title"], level=2)
    sec10_heading.runs[0].font.name = DOCX_FONT

    validation_label = doc.add_paragraph("Validation Checks:")
    validation_label.runs[0].font.name = DOCX_FONT
    validation_label.runs[0].font.bold = True
    add_docx_bullets(doc, report_content["section10"]["validation_points"])

    metrics_label = doc.add_paragraph("Quality Indicators:")
    metrics_label.runs[0].font.name = DOCX_FONT
    metrics_label.runs[0].font.bold = True
    add_docx_bullets(doc, report_content["section10"]["quality_metrics"])

    doc.add_page_break()

    # Section 11
    sec11_heading = doc.add_heading(report_content["section11"]["title"], level=2)
    sec11_heading.runs[0].font.name = DOCX_FONT

    status_para = doc.add_paragraph(f"Status: {report_content['section11']['status']}")
    status_para.runs[0].font.name = DOCX_FONT
    status_para.runs[0].font.bold = True

    for result in report_content["section11"]["results"]:
        result_para = doc.add_paragraph(result, style='List Bullet')
        for run in result_para.runs:
            run.font.name = DOCX_FONT
            run.font.size = Pt(10)

    # Section 12
    sec12_heading = doc.add_heading(report_content["section12"]["title"], level=2)
    sec12_heading.runs[0].font.name = DOCX_FONT

    contrib_label = doc.add_paragraph("Key Contributions:")
    contrib_label.runs[0].font.name = DOCX_FONT
    contrib_label.runs[0].font.bold = True

    for contrib in report_content["section12"]["contributions"]:
        contrib_para = doc.add_paragraph(contrib, style='List Bullet')
        for run in contrib_para.runs:
            run.font.name = DOCX_FONT
            run.font.size = Pt(10)

    future_label = doc.add_paragraph("\nFuture Enhancements:")
    future_label.runs[0].font.name = DOCX_FONT
    future_label.runs[0].font.bold = True

    for future in report_content["section12"]["future"]:
        future_para = doc.add_paragraph(future, style='List Bullet')
        for run in future_para.runs:
            run.font.name = DOCX_FONT
            run.font.size = Pt(10)

    status_final = doc.add_paragraph("\nProject Status: ✅ Successfully Completed and Production-Ready")
    status_final.runs[0].font.name = DOCX_FONT
    status_final.runs[0].font.bold = True
    
    # Save
    doc.save("CVCreationUsingLLM_Project_Report.docx")
    print("✅ DOCX report generated: CVCreationUsingLLM_Project_Report.docx")
    print(f"   Font used: {DOCX_FONT}")

# ============ PDF Generation ============
def generate_pdf():
    pdf_path = "CVCreationUsingLLM_Project_Report.pdf"
    doc = SimpleDocTemplate(pdf_path, pagesize=letter,
                           topMargin=0.5*inch, bottomMargin=0.5*inch,
                           leftMargin=0.75*inch, rightMargin=0.75*inch)
    
    story = []
    styles = getSampleStyleSheet()
    
    # Custom styles with cross-platform font
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontName=PDF_FONT,
        fontSize=16,
        textColor=RGBColor(0, 51, 102),
        spaceAfter=12,
        alignment=1  # CENTER alignment
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontName=PDF_FONT,
        fontSize=12,
        textColor=RGBColor(0, 51, 102),
        spaceAfter=6,
        spaceBefore=6
    )
    
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['Normal'],
        fontName=PDF_FONT,
        fontSize=9.5,
        spaceAfter=6
    )

    mono_style = ParagraphStyle(
        'MonospaceBody',
        parent=styles['Code'],
        fontName='Courier',
        fontSize=7.6,
        leading=8.4,
        leftIndent=6,
        rightIndent=6,
        spaceAfter=6
    )
    
    # Title
    story.append(Paragraph(report_content["title"].replace("\n", "<br/>"), title_style))
    story.append(Spacer(1, 0.2*inch))
    
    # Abstract
    story.append(Paragraph("Abstract", heading_style))
    story.append(Paragraph(report_content["abstract"], body_style))
    story.append(Spacer(1, 0.1*inch))
    
    # Sections 1-2
    for i in [1, 2]:
        section_key = f"section{i}"
        story.append(Paragraph(report_content[section_key]["title"], heading_style))
        story.append(Paragraph(report_content[section_key]["content"], body_style))
        story.append(Spacer(1, 0.1*inch))
    
    # Section 3
    story.append(Paragraph(report_content["section3"]["title"], heading_style))
    add_pdf_bullets(story, report_content["section3"]["objectives"], body_style)
    story.append(Spacer(1, 0.1*inch))
    
    # Section 4
    story.append(Paragraph(report_content["section4"]["title"], heading_style))
    story.append(Paragraph("<b>Technology Stack:</b>", body_style))
    for tech in report_content["section4"]["tech_stack"]:
        story.append(Paragraph(f"<bullet>•</bullet> {tech}", body_style))

    story.append(Paragraph("<b>Core Modules:</b>", body_style))
    for module, responsibility, feature in report_content["section4"]["modules"]:
        story.append(Paragraph(f"<bullet>•</bullet> {module}: {responsibility} ({feature})", body_style))

    story.append(PageBreak())

    # Section 5 - Architecture
    story.append(Paragraph(report_content["section5"]["title"], heading_style))
    story.append(Paragraph(report_content["section5"]["overview"], body_style))
    story.append(Paragraph("<b>Architecture Diagram:</b>", body_style))
    story.append(Preformatted("\n".join(report_content["section5"]["diagram"]), mono_style))
    story.append(Paragraph("<b>Component Interaction Notes:</b>", body_style))
    add_pdf_bullets(story, report_content["section5"]["flow"], body_style)

    story.append(PageBreak())

    # Sections 6-9 Trade-off tables
    for section_key in ["section6", "section7", "section8", "section9"]:
        story.append(Paragraph(report_content[section_key]["title"], heading_style))
        add_pdf_tradeoff_table(story, report_content[section_key], body_style)

    # Section 7b – Chunking Implementation Details
    sec7b = report_content["section7b"]
    story.append(Paragraph(sec7b["title"], heading_style))
    story.append(Paragraph("<b>Threshold Configuration:</b>", body_style))

    thresh_table_data = [["Module", "Doc Threshold", "Chunk Size", "Overlap"]]
    for row_data in sec7b["thresholds"]:
        thresh_table_data.append(list(row_data))
    thresh_tbl = Table(thresh_table_data, colWidths=[1.6*inch, 1.6*inch, 1.6*inch, 1.6*inch])
    thresh_tbl.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#D9E2F3')),
        ('FONTNAME', (0, 0), (-1, 0), PDF_FONT + '-Bold' if False else PDF_FONT),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
    ]))
    story.append(thresh_tbl)
    story.append(Spacer(1, 0.08*inch))

    story.append(Paragraph("<b>Chunking Stages:</b>", body_style))
    add_pdf_bullets(story, sec7b["stages"], body_style)

    # Section 7c – Keyword Optimisation
    story.append(Paragraph(report_content["section7c"]["title"], heading_style))
    add_pdf_tradeoff_table(story, report_content["section7c"], body_style)

    story.append(PageBreak())

    # Section 10
    story.append(Paragraph(report_content["section10"]["title"], heading_style))
    story.append(Paragraph("<b>Validation Checks:</b>", body_style))
    add_pdf_bullets(story, report_content["section10"]["validation_points"], body_style)
    story.append(Paragraph("<b>Quality Indicators:</b>", body_style))
    add_pdf_bullets(story, report_content["section10"]["quality_metrics"], body_style)

    story.append(Spacer(1, 0.05*inch))
    
    # Section 11
    story.append(Paragraph(report_content["section11"]["title"], heading_style))
    story.append(Paragraph(f"<b>Status:</b> {report_content['section11']['status']}", body_style))
    add_pdf_bullets(story, report_content["section11"]["results"], body_style)
    story.append(Spacer(1, 0.1*inch))
    
    # Section 12
    story.append(Paragraph(report_content["section12"]["title"], heading_style))
    story.append(Paragraph("<b>Key Contributions:</b>", body_style))
    add_pdf_bullets(story, report_content["section12"]["contributions"], body_style)

    story.append(Paragraph("<b>Future Enhancements:</b>", body_style))
    add_pdf_bullets(story, report_content["section12"]["future"], body_style)

    story.append(Spacer(1, 0.05*inch))
    story.append(Paragraph("<b>Project Status:</b> ✅ Successfully Completed", body_style))
    
    # Build PDF
    doc.build(story)
    print("✅ PDF report generated: CVCreationUsingLLM_Project_Report.pdf")
    print(f"   Font used: {PDF_FONT}")

# ============ Main Execution ============
if __name__ == "__main__":
    print("=" * 60)
    print("CV CREATION PROJECT REPORT GENERATOR")
    print("=" * 60)
    print("\nCross-Platform Font Configuration:")
    print(f"  • DOCX Font: {DOCX_FONT} (Windows, Mac, Linux compatible)")
    print(f"  • PDF Font: {PDF_FONT} (Universal cross-platform font)")
    print("\n" + "=" * 60)
    print("\nGenerating reports...\n")
    
    try:
        generate_docx()
        generate_pdf()
        print("\n" + "=" * 60)
        print("✅ SUCCESS! Both reports generated successfully!")
        print("=" * 60)
        print("\nFiles created in Summary folder:")
        print("  1. CVCreationUsingLLM_Project_Report.pdf")
        print("  2. CVCreationUsingLLM_Project_Report.docx")
        print("\n✓ Reports are expanded with architecture and trade-off analysis")
        print("✓ All fonts are cross-platform compatible")
        print("✓ Reports are production-ready")
    except Exception as e:
        print(f"\n❌ Error: {str(e)}")
        import traceback
        traceback.print_exc()
